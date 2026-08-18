from __future__ import annotations

import importlib.util
import os
import re
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Any

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
LOCAL_STYLE_PATH = Path(__file__).with_name("report_style.py")
STYLE_PATH = LOCAL_STYLE_PATH if LOCAL_STYLE_PATH.exists() else ROOT / "work" / "build_report.py"
_spec = importlib.util.spec_from_file_location("report_style", STYLE_PATH)
style = importlib.util.module_from_spec(_spec)
assert _spec and _spec.loader
_spec.loader.exec_module(style)


class ReportError(ValueError):
    pass


def _clean(value: str) -> str:
    return re.sub(r"\s+", "", value or "").replace("％", "%")


def _number(value: str, field: str) -> float:
    text = _clean(value).replace("%", "").replace(",", "")
    try:
        return float(text)
    except ValueError as exc:
        raise ReportError(f"{field}无法解析：{value}") from exc


def _integer(value: str, field: str) -> int:
    return int(round(_number(value, field)))


def _infer_age(filename: str, paragraphs: list[str]) -> str:
    sources = [Path(filename).stem] + paragraphs[:12]
    patterns = [
        r"([一二三四五六七八九十0-9]{1,4})年级",
        r"(小学|初中|高中)(?:低|中|高)?年级段?",
    ]
    for source in sources:
        for pattern in patterns:
            match = re.search(pattern, source)
            if match:
                value = match.group(0)
                return value if value.endswith("年级") else value
    return ""


def parse_docx(content: bytes, filename: str) -> dict[str, Any]:
    if not filename.lower().endswith(".docx"):
        raise ReportError("仅支持 .docx 格式的测评报告")
    try:
        doc = Document(BytesIO(content))
    except Exception as exc:
        raise ReportError("文档无法打开，可能已损坏或不是有效的 Word 文件") from exc

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables: dict[str, list[list[str]]] = {}
    expected = {
        "整体": ["有效样本", "平均分", "中位数", "标准差", "最低分", "最高分"],
        "分布": ["分数段", "人数", "占比"],
        "一级": ["一级知识维度", "题目数", "满分", "得分率"],
        "二级": ["二级知识维度", "题目数", "满分", "得分率"],
        "题目": ["题号", "一级知识维度", "二级知识维度", "满分", "得分率"],
    }
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        header = [_clean(x) for x in rows[0]]
        for key, wanted in expected.items():
            if all(_clean(x) in header for x in wanted):
                tables[key] = rows
                break
    missing = [key for key in expected if key not in tables]
    if missing:
        raise ReportError("缺少或无法识别以下数据表：" + "、".join(missing))

    overview = tables["整体"]
    if len(overview) < 2:
        raise ReportError("整体成绩总览缺少数据行")
    values = overview[1]
    dataset: dict[str, Any] = {
        "source": filename,
        "name": _infer_age(filename, paragraphs),
        "n": _integer(values[0], "有效样本"),
        "mean": _number(values[1], "平均分"),
        "median": _number(values[2], "中位数"),
        "sd": _number(values[3], "标准差"),
        "min": _number(values[4], "最低分"),
        "max": _number(values[5], "最高分"),
    }
    dataset["dist"] = [
        [row[0].replace("%", "分"), _integer(row[1], "分数段人数"), _number(row[2], "分数段占比")]
        for row in tables["分布"][1:] if row and _clean(row[0])
    ]
    dataset["primary"] = [
        [row[0], _integer(row[1], "一级维度题目数"), _number(row[2], "一级维度满分"), _number(row[3], "一级维度得分率")]
        for row in tables["一级"][1:] if row and _clean(row[0])
    ]
    dataset["secondary"] = [
        [row[0], _integer(row[1], "二级维度题目数"), _number(row[2], "二级维度满分"), _number(row[3], "二级维度得分率")]
        for row in tables["二级"][1:] if row and _clean(row[0])
    ]
    dataset["items"] = [
        [_integer(row[0], "题号"), row[1], row[2], _number(row[3], "题目满分"), _number(row[4], "题目得分率")]
        for row in tables["题目"][1:] if row and _clean(row[0])
    ]
    if not dataset["dist"] or not dataset["primary"] or not dataset["secondary"]:
        raise ReportError("成绩分布或知识维度数据为空")
    return dataset


def _pass_stats(dataset: dict[str, Any]) -> tuple[int, float]:
    below = 0
    for label, count, _ in dataset["dist"]:
        if "以下" in label and ("60" in label or "六十" in label):
            below += int(count)
    passed = int(dataset["n"]) - below
    return passed, passed / int(dataset["n"]) * 100 if dataset["n"] else 0


def _excellent_count(dataset: dict[str, Any]) -> int:
    return sum(int(count) for label, count, _ in dataset["dist"] if "90" in label and ("以上" in label or "及" in label))


def validate_request(school: str, year: Any, datasets: list[dict[str, Any]]) -> tuple[str, int]:
    school = re.sub(r"[\\/:*?\"<>|]", "", (school or "").strip())
    if not school:
        raise ReportError("请填写学校名称")
    try:
        year = int(year)
    except (TypeError, ValueError) as exc:
        raise ReportError("报告年度格式不正确") from exc
    if year < 2000 or year > 2100:
        raise ReportError("报告年度应在 2000 至 2100 之间")
    if len(datasets) < 2:
        raise ReportError("至少需要两份有效的年龄段报告")
    names = [str(d.get("name", "")).strip() for d in datasets]
    if any(not name for name in names):
        raise ReportError("存在未识别的年龄段，请先补充年龄段名称")
    if len(set(names)) != len(names):
        raise ReportError("年龄段名称重复，请检查后再生成")
    return school, year


def _fonts():
    regular_candidates = [
        Path(value) for value in [
            os.environ.get("REPORT_FONT_REGULAR", ""),
            r"C:\Windows\Fonts\msyh.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf",
        ] if value
    ]
    bold_candidates = [
        Path(value) for value in [
            os.environ.get("REPORT_FONT_BOLD", ""),
            r"C:\Windows\Fonts\msyhbd.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Bold.otf",
        ] if value
    ]
    regular_path = next((path for path in regular_candidates if path.is_file()), None)
    bold_path = next((path for path in bold_candidates if path.is_file()), None)
    if regular_path is None or bold_path is None:
        raise ReportError("服务器缺少中文字体，暂时无法生成图表")
    return (
        lambda size: ImageFont.truetype(str(regular_path), size),
        lambda size: ImageFont.truetype(str(bold_path), size),
    )


def _make_charts(datasets: list[dict[str, Any]], chart_dir: Path) -> tuple[Path, Path, list[Path]]:
    chart_dir.mkdir(parents=True, exist_ok=True)
    font, bold = _fonts()
    colors = ["#2E75B6", "#D6A126", "#2A7F7F", "#C65911"]
    ink, grid = "#30343A", "#E4E8ED"

    def canvas(height=720):
        image = Image.new("RGB", (1640, height), "white")
        return image, ImageDraw.Draw(image)

    image, draw = canvas()
    draw.text((70, 20), "各年龄段平均分与达标率", font=bold(44), fill=ink)
    base_y, top_y = 610, 110
    for value in range(0, 101, 20):
        y = base_y - int((base_y - top_y) * value / 100)
        draw.line((105, y, 1560, y), fill=grid, width=2)
        draw.text((30, y - 16), str(value), font=font(25), fill="#777777")
    group_width = 1300 / len(datasets)
    for index, dataset in enumerate(datasets):
        center = 190 + group_width * (index + 0.5)
        _, pass_rate = _pass_stats(dataset)
        for offset, value, color in [(-55, dataset["mean"], colors[index % len(colors)]), (15, pass_rate, "#7B8794")]:
            height = int((base_y - top_y) * value / 100)
            x = int(center + offset)
            draw.rounded_rectangle((x, base_y - height, x + 58, base_y), radius=7, fill=color)
            draw.text((x + 29, base_y - height - 32), f"{value:.1f}", font=bold(24), fill=ink, anchor="mm")
        draw.text((center, 650), dataset["name"], font=bold(27), fill=ink, anchor="mm")
    draw.rectangle((1120, 29, 1150, 57), fill=colors[0])
    draw.text((1164, 25), "平均分", font=font(25), fill=ink)
    draw.rectangle((1340, 29, 1370, 57), fill="#7B8794")
    draw.text((1384, 25), "达标率", font=font(25), fill=ink)
    overall_path = chart_dir / "overall.png"
    image.save(overall_path)

    labels = [row[0] for row in datasets[0]["dist"]]
    image, draw = canvas()
    draw.text((70, 20), "各年龄段成绩分布", font=bold(44), fill=ink)
    base_y, top_y = 610, 110
    max_rate = max(row[2] for data in datasets for row in data["dist"])
    axis_max = max(50, int((max_rate + 9) // 10 * 10))
    for value in range(0, axis_max + 1, 10):
        y = base_y - int((base_y - top_y) * value / axis_max)
        draw.line((100, y, 1580, y), fill=grid, width=2)
        draw.text((20, y - 15), f"{value}%", font=font(23), fill="#777777")
    group_width = 1440 / max(1, len(labels))
    bar_width = min(48, int(group_width / (len(datasets) + 1)))
    for li, label in enumerate(labels):
        center = 110 + group_width * (li + 0.5)
        for di, dataset in enumerate(datasets):
            match = next((row for row in dataset["dist"] if row[0] == label), None)
            value = match[2] if match else 0
            x = int(center - len(datasets) * bar_width / 2 + di * bar_width)
            height = int((base_y - top_y) * value / axis_max)
            draw.rectangle((x, base_y - height, x + bar_width - 5, base_y), fill=colors[di % len(colors)])
            if len(datasets) <= 2:
                draw.text((x + (bar_width - 5) / 2, base_y - height - 25), f"{value:.1f}%", font=font(20), fill=ink, anchor="mm")
        draw.text((center, 650), label, font=font(22), fill=ink, anchor="mm")
    for index, dataset in enumerate(datasets):
        x = 930 + index * 160
        draw.rectangle((x, 35, x + 22, 56), fill=colors[index % len(colors)])
        draw.text((x + 30, 29), dataset["name"], font=font(23), fill=ink)
    distribution_path = chart_dir / "distribution.png"
    image.save(distribution_path)

    primary_paths = []
    for index, dataset in enumerate(datasets):
        rows = sorted(dataset["primary"], key=lambda item: item[3])
        height = max(650, 180 + len(rows) * 92)
        image, draw = canvas(height)
        draw.text((60, 20), f"{dataset['name']}一级知识维度得分率", font=bold(42), fill=ink)
        left, right, top, row_height = 650, 1490, 105, 90
        ref_x = left + int((right - left) * 0.6)
        draw.line((ref_x, top - 8, ref_x, top + row_height * len(rows) - 20), fill="#C65911", width=3)
        draw.text((ref_x + 8, top - 40), "60%参考线", font=font(23), fill="#C65911")
        for ri, (label, _, _, value) in enumerate(rows):
            y = top + ri * row_height
            draw.text((620, y + 24), label, font=font(28), fill=ink, anchor="ra")
            draw.rounded_rectangle((left, y + 8, right, y + 54), radius=7, fill="#EDF0F4")
            bar_right = left + int((right - left) * value / 100)
            draw.rounded_rectangle((left, y + 8, bar_right, y + 54), radius=7, fill=colors[index % len(colors)])
            draw.text((bar_right + 12, y + 31), f"{value:.2f}%", font=bold(25), fill=ink, anchor="lm")
        path = chart_dir / f"primary_{index}.png"
        image.save(path)
        primary_paths.append(path)
    return overall_path, distribution_path, primary_paths


def _add_cover(doc, school: str, year: int):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.set_paragraph(p, after=8)
    style.add_text(p, f"{year}年", size=11, bold=True, color=style.GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.set_paragraph(p, after=10, line=1.1)
    style.add_text(p, school, size=24, bold=True, color=style.NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.set_paragraph(p, after=8)
    style.add_text(p, "学生素养测评数据分析报告", size=29, bold=True, color=style.NAVY)
    doc.add_page_break()


def _add_overview(doc, datasets: list[dict[str, Any]], overall_chart: Path, distribution_chart: Path):
    total = sum(int(d["n"]) for d in datasets)
    mean = sum(float(d["mean"]) * int(d["n"]) for d in datasets) / total
    passes = [_pass_stats(d) for d in datasets]
    passed = sum(item[0] for item in passes)
    pass_rate = passed / total * 100
    excellent = sum(_excellent_count(d) for d in datasets)

    style.add_heading(doc, "一、学校整体情况", 1)
    details = "，".join(f"{d['name']}达标率为{passes[i][1]:.2f}%" for i, d in enumerate(datasets))
    style.add_callout(doc, "整体表现", f"本次测评共纳入{total}份有效记录，学校整体平均分为{mean:.2f}分，达标{passed}人，达标率为{pass_rate:.2f}%。{details}。", fill=style.PALE_GOLD, accent=style.RED)

    style.add_heading(doc, "1. 整体成绩概况", 2)
    style.add_body(doc, f"从整体成绩看，学校共有{passed}名学生达到60分及以上。各年龄段在平均分、达标率和高分段人数方面呈现出不同特点。")
    headers = ["指标", "学校整体"] + [d["name"] for d in datasets]
    rows = [
        ["有效样本", str(total)] + [str(d["n"]) for d in datasets],
        ["平均分", f"{mean:.2f}"] + [f"{d['mean']:.2f}" for d in datasets],
        ["达标人数（≥60分）", str(passed)] + [str(x[0]) for x in passes],
        ["达标率", f"{pass_rate:.2f}%"] + [f"{x[1]:.2f}%" for x in passes],
        ["优秀段人数（≥90分）", str(excellent)] + [str(_excellent_count(d)) for d in datasets],
    ]
    col_count = len(headers)
    widths = [2500] + [int(6860 / (col_count - 1))] * (col_count - 1)
    style.add_table(doc, headers, rows, widths, font=8.8)
    style.add_picture(doc, overall_chart, width=6.35, caption="图1  各年龄段平均分与达标率对比")

    doc.add_page_break()
    style.add_heading(doc, "2. 成绩分布与年龄段特点", 2)
    labels = []
    for dataset in datasets:
        for label, _, _ in dataset["dist"]:
            if label not in labels:
                labels.append(label)
    distribution_rows = []
    for label in labels:
        counts = [next((int(r[1]) for r in d["dist"] if r[0] == label), 0) for d in datasets]
        count = sum(counts)
        distribution_rows.append([label, count, f"{count / total * 100:.2f}%"] + counts)
    headers = ["分数段", "学校人数", "学校占比"] + [d["name"] for d in datasets]
    widths = [1900, 1300, 1500] + [int(4660 / len(datasets))] * len(datasets)
    distribution_table = style.add_table(doc, headers, distribution_rows, widths, font=8.5)
    header_properties = distribution_table.rows[0]._tr.get_or_add_trPr()
    for marker in list(header_properties.findall(qn("w:tblHeader"))):
        header_properties.remove(marker)
    for row in distribution_table.rows:
        row_properties = row._tr.get_or_add_trPr()
        if row_properties.find(qn("w:cantSplit")) is None:
            row_properties.append(style.OxmlElement("w:cantSplit"))
    style.add_picture(doc, distribution_chart, caption="图2  各年龄段成绩分布对比")
    below = total - passed
    style.add_bullet(doc, f"学校整体：60分以下{below}人，占{below / total * 100:.2f}%；60分及以上{passed}人，占{pass_rate:.2f}%。")
    for index, dataset in enumerate(datasets):
        high = sum(int(r[1]) for r in dataset["dist"] if any(mark in r[0] for mark in ["80", "90"]))
        style.add_bullet(doc, f"{dataset['name']}：达标{passes[index][0]}人，达标率为{passes[index][1]:.2f}%；80分及以上{high}人。")

    style.add_heading(doc, "3. 各能力维度主要表现", 2)
    style.add_body(doc, "从一级知识维度看，各年龄段均呈现出相对优势与需要关注的方向。以下结果用于概括各年龄段内部的主要特点。")
    dimension_rows = []
    for dataset in datasets:
        ordered = sorted(dataset["primary"], key=lambda item: item[3], reverse=True)
        dimension_rows.append([dataset["name"], f"{ordered[0][0]}（{ordered[0][3]:.2f}%）", f"{ordered[-1][0]}（{ordered[-1][3]:.2f}%）"])
    style.add_table(doc, ["年龄段", "相对优势维度", "需要关注的维度"], dimension_rows, [1700, 3830, 3830], font=9.0, aligns=[WD_ALIGN_PARAGRAPH.LEFT] * 3)


def _add_age_section(doc, dataset: dict[str, Any], part_no: str, chart_path: Path):
    name = dataset["name"]
    passed, pass_rate = _pass_stats(dataset)
    primary_sorted = sorted(dataset["primary"], key=lambda item: item[3], reverse=True)
    secondary_sorted = sorted(dataset["secondary"], key=lambda item: item[3], reverse=True)
    below_row = next((row for row in dataset["dist"] if "以下" in row[0] and "60" in row[0]), dataset["dist"][0])

    style.add_heading(doc, f"{part_no}、{name}测评情况", 1)
    style.add_heading(doc, "1. 关键结论摘要", 2)
    style.add_bullet(doc, f"本次纳入有效记录{dataset['n']}份，平均分为{dataset['mean']:.2f}分，达标{passed}人，达标率为{pass_rate:.2f}%。")
    style.add_bullet(doc, f"{below_row[0]}共{below_row[1]}人，占{below_row[2]:.2f}%。")
    style.add_bullet(doc, f"知识表现中，{primary_sorted[0][0]}相对较高，{primary_sorted[-1][0]}相对较低，二者得分率相差{primary_sorted[0][3] - primary_sorted[-1][3]:.2f}个百分点。")

    style.add_heading(doc, "2. 数据口径与样本概况", 2)
    style.add_body(doc, f"本次统计纳入{dataset['n']}份有效记录，测评成绩满分为100分。")
    style.add_heading(doc, "3. 整体表现分析", 2)
    style.add_body(doc, f"平均分为{dataset['mean']:.2f}分，中位数为{dataset['median']:.2f}分，标准差为{dataset['sd']:.2f}分。")
    style.add_table(doc, ["有效样本", "平均分", "中位数", "标准差", "最低分", "最高分"], [[dataset["n"], f"{dataset['mean']:.2f}", f"{dataset['median']:.2f}", f"{dataset['sd']:.2f}", f"{dataset['min']:.2f}", f"{dataset['max']:.2f}"]], [1560] * 6, font=9.2)
    style.add_table(doc, ["分数段", "人数", "占比"], [[a, b, f"{c:.2f}%"] for a, b, c in dataset["dist"]], [4200, 2400, 2760], font=9.5)

    style.add_heading(doc, "4. 一级知识维度分析", 2)
    primary_avg = sum(item[3] for item in dataset["primary"]) / len(dataset["primary"])
    style.add_body(doc, f"本板块共覆盖{len(dataset['primary'])}个维度，平均得分率为{primary_avg:.2f}%。{primary_sorted[0][0]}相对较高（{primary_sorted[0][3]:.2f}%），{primary_sorted[-1][0]}相对较低（{primary_sorted[-1][3]:.2f}%）。")
    style.add_picture(doc, chart_path, caption=f"图  {name}一级知识维度得分率")
    style.add_table(doc, ["一级知识维度", "题目数", "满分", "得分率"], [[n, q, f"{m:.2f}", f"{r:.2f}%"] for n, q, m, r in dataset["primary"]], [4800, 1320, 1320, 1920], font=8.9)

    style.add_heading(doc, "5. 二级知识维度分析", 2)
    secondary_avg = sum(item[3] for item in dataset["secondary"]) / len(dataset["secondary"])
    style.add_body(doc, f"本板块共覆盖{len(dataset['secondary'])}个维度，平均得分率为{secondary_avg:.2f}%。{secondary_sorted[0][0]}相对较高（{secondary_sorted[0][3]:.2f}%），{secondary_sorted[-1][0]}相对较低（{secondary_sorted[-1][3]:.2f}%）。")
    half = (len(dataset["secondary"]) + 1) // 2
    rows = []
    for index in range(half):
        left = dataset["secondary"][index]
        right = dataset["secondary"][index + half] if index + half < len(dataset["secondary"]) else ["", 0, 0, 0]
        rows.append([left[0], f"{left[3]:.2f}%", right[0], f"{right[3]:.2f}%" if right[0] else ""])
    style.add_table(doc, ["二级知识维度", "得分率", "二级知识维度", "得分率"], rows, [3300, 1380, 3300, 1380], font=8.3)

    style.add_heading(doc, "6. 重点题目诊断", 2)
    items = sorted(dataset["items"], key=lambda item: item[4])[:5]
    average = sum(item[4] for item in items) / len(items)
    style.add_body(doc, f"按相对得分率选取{len(items)}道重点复核题，平均得分率为{average:.2f}%。建议结合题干语言、选项干扰、作答方式和学生认知负荷逐题复核。")
    style.add_table(doc, ["题号", "一级知识维度", "二级知识维度", "得分率"], [[x[0], x[1], x[2], f"{x[4]:.2f}%"] for x in items], [900, 3300, 3300, 1860], font=8.8)


def build_report(school: str, year: Any, datasets: list[dict[str, Any]], output_path: Path) -> Path:
    school, year = validate_request(school, year, datasets)
    with tempfile.TemporaryDirectory(prefix="school-report-") as temp:
        chart_dir = Path(temp) / "charts"
        overall, distribution, primary = _make_charts(datasets, chart_dir)
        doc = style.setup_document()
        _add_cover(doc, school, year)
        _add_overview(doc, datasets, overall, distribution)
        numerals = ["二", "三", "四", "五"]
        for index, dataset in enumerate(datasets):
            _add_age_section(doc, dataset, numerals[index], primary[index])
        properties = doc.core_properties
        properties.title = f"{school}学生素养测评数据分析报告"
        properties.subject = "学生素养测评数据分析"
        properties.author = ""
        properties.keywords = "素养测评, 数据分析"
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
    return output_path
