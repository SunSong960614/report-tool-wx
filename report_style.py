from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\PC\Documents\Codex\2026-08-18\new-chat")
WORK = ROOT / "work"
OUT = ROOT / "outputs"
CHARTS = WORK / "charts"
OUT.mkdir(parents=True, exist_ok=True)
CHARTS.mkdir(parents=True, exist_ok=True)
OUTPUT = OUT / "XXX学校学生素养测评数据分析报告.docx"

NAVY = "17365D"
BLUE = "2E75B6"
CYAN = "5B9BD5"
TEAL = "2A7F7F"
GREEN = "70AD47"
GOLD = "D6A126"
RED = "C65911"
PALE_BLUE = "EAF2F8"
PALE_GOLD = "FFF4D6"
LIGHT = "F4F6F8"
MID = "D9E2F3"
GRAY = "666666"
BLACK = "222222"
WHITE = "FFFFFF"


data_34 = {
    "name": "三四年级",
    "n": 30, "mean": 64.53, "median": 68.00, "sd": 23.13, "min": 8.00, "max": 96.00,
    "dist": [("60分以下", 11, 36.67), ("60–69分", 5, 16.67), ("70–79分", 4, 13.33), ("80–89分", 6, 20.00), ("90分及以上", 4, 13.33)],
    "primary": [
        ("AI伦理、安全与社会责任", 7, 28, 51.90),
        ("AI工具应用与人机协作", 5, 20, 71.33),
        ("AI关键技术与应用形态", 4, 16, 55.83),
        ("数据、算法与模型", 5, 20, 72.67),
        ("AI基础认知与智能社会", 2, 8, 58.33),
        ("AI系统设计与工程实践", 3, 12, 63.33),
    ],
    "secondary": [
        ("环境影响与可持续性", 46.67), ("AI输出核验与纠错", 80.00), ("计算机视觉", 73.33),
        ("AI工具识别与选择", 66.67), ("模型训练与推理", 76.67), ("任务表达与提示设计", 83.33),
        ("规则遵守与责任归属", 53.33), ("数据分类与标注", 60.00), ("AI能力边界与不确定性", 76.67),
        ("机器学习基本概念", 76.67), ("包容、公平与可访问性", 43.33), ("监督学习与常见任务", 80.00),
        ("人机分工与协同决策", 56.67), ("公平、偏见与非歧视", 40.00), ("AI定义与基本特征", 40.00),
        ("深度伪造与智能诈骗", 76.67), ("模型训练、测试与优化流程", 53.33), ("效果评估与反馈迭代", 53.33),
        ("个人信息与数据隐私", 43.33), ("语音识别与语音合成", 43.33), ("问题界定与需求分析", 83.33),
        ("学术诚信与规范使用", 60.00), ("数据整理与清洗", 70.00), ("传感器与环境感知", 33.33),
        ("科学与领域知识迁移", 70.00), ("自然语言处理与机器翻译", 73.33),
    ],
    "items": [(24, "AI关键技术与应用形态", "传感器与环境感知", 33.33),
              (14, "AI伦理、安全与社会责任", "公平、偏见与非歧视", 40.00),
              (15, "AI基础认知与智能社会", "AI定义与基本特征", 40.00),
              (11, "AI伦理、安全与社会责任", "包容、公平与可访问性", 43.33),
              (19, "AI伦理、安全与社会责任", "个人信息与数据隐私", 43.33)],
}

data_56 = {
    "name": "五六年级",
    "n": 30, "mean": 56.00, "median": 60.00, "sd": 20.29, "min": 16.00, "max": 88.00,
    "dist": [("60分以下", 14, 46.67), ("60–69分", 8, 26.67), ("70–79分", 3, 10.00), ("80–89分", 5, 16.67), ("90分及以上", 0, 0.00)],
    "primary": [
        ("AI伦理、安全与社会责任", 5, 20, 54.67),
        ("AI工具应用与人机协作", 2, 8, 50.00),
        ("AI基础认知与智能社会", 5, 20, 64.67),
        ("数据、算法与模型", 6, 24, 62.78),
        ("AI关键技术与应用形态", 3, 12, 41.11),
        ("AI系统设计与工程实践", 4, 16, 50.83),
    ],
    "secondary": [
        ("技术自主与公民责任", 80.00), ("工具结果评价与迭代优化", 43.33), ("AI发展与典型应用场景", 66.67),
        ("数据特征与表示", 40.00), ("数据来源与采集", 86.67), ("AI社会影响与职业变化", 66.67),
        ("透明度与可解释性", 50.00), ("模型评价指标", 63.33), ("推荐系统", 33.33),
        ("生成式人工智能", 40.00), ("任务拆解与流程规划", 43.33), ("算法与规则", 73.33),
        ("生物特征信息保护", 40.00), ("技术适配与方案选择", 43.33), ("跨学科问题分析", 56.67),
        ("账号安全与数字身份", 46.67), ("AI与一般自动化的区别", 60.00), ("人类智能与机器智能差异", 50.00),
        ("智能体与智能机器人", 50.00), ("数据质量与多样性", 50.00), ("可靠性与鲁棒性", 43.33),
        ("数据偏差与模型偏差", 63.33), ("知识产权与AI生成标注", 56.67), ("AI系统组成与数据流", 73.33),
        ("人类主体性与决策责任", 80.00),
    ],
    "items": [(9, "AI关键技术与应用形态", "推荐系统", 33.33),
              (4, "数据、算法与模型", "数据特征与表示", 40.00),
              (10, "AI关键技术与应用形态", "生成式人工智能", 40.00),
              (13, "AI伦理、安全与社会责任", "生物特征信息保护", 40.00),
              (2, "AI工具应用与人机协作", "工具结果评价与迭代优化", 43.33)],
}


def level(rate):
    if rate >= 80: return "优势"
    if rate >= 70: return "良好"
    if rate >= 60: return "达标"
    return "待提升"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(table):
    repeat_header(table.rows[0])


def set_run(run, size=10.5, bold=False, color=BLACK, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = rgb(color)


def clear_paragraph(p):
    for child in list(p._p):
        if child.tag != qn("w:pPr"):
            p._p.remove(child)


def add_text(p, text, size=10.5, bold=False, color=BLACK):
    r = p.add_run(str(text))
    set_run(r, size=size, bold=bold, color=color)
    return r


def set_paragraph(p, before=0, after=6, line=1.25, align=None, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None: p.alignment = align
    if keep: pf.keep_with_next = True


def add_heading(doc, text, level_num=1):
    p = doc.add_paragraph(style=f"Heading {level_num}")
    add_text(p, text, size={1:16, 2:13, 3:11.5}[level_num], bold=True,
             color={1:NAVY, 2:BLUE, 3:TEAL}[level_num])
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True)
        add_text(p, text[len(bold_prefix):])
    else:
        add_text(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    add_text(p, text)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    t = doc.add_table(rows=1, cols=1)
    set_table_geometry(t, [9360])
    cell = t.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph(p, after=0, line=1.25)
    add_text(p, label + "  ", size=10.5, bold=True, color=accent)
    add_text(p, text, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def add_table(doc, headers, rows, widths, font=9.5, aligns=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_table_geometry(t, widths)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, NAVY)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        set_paragraph(p, after=0, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, h, size=font, bold=True, color=WHITE)
    for r_idx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            if r_idx % 2 == 1: set_cell_shading(c, LIGHT)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = c.paragraphs[0]
            align = aligns[i] if aligns else (WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
            set_paragraph(p, after=0, line=1.08, align=align)
            color = RED if str(val) == "待提升" else BLACK
            add_text(p, val, size=font, bold=(i == 0), color=color)
    set_repeat_table_header(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def add_picture(doc, path, width=6.35, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(cp, after=7, line=1.0)
        add_text(cp, caption, size=8.5, color=GRAY)


def make_charts():
    font_path = r"C:\Windows\Fonts\msyh.ttc"
    bold_path = r"C:\Windows\Fonts\msyhbd.ttc"
    font = lambda n: ImageFont.truetype(font_path, n)
    bold = lambda n: ImageFont.truetype(bold_path, n)
    blue, gold, grid, ink, red = "#2E75B6", "#D6A126", "#E4E8ED", "#30343A", "#C65911"

    def canvas():
        im = Image.new("RGB", (1640, 720), "white")
        return im, ImageDraw.Draw(im)

    # Overall grouped bars.
    im, d = canvas(); base_y, top_y = 610, 90
    d.text((85, 30), "平均分与达标率对比", font=bold(34), fill=ink)
    for pct in range(0, 81, 20):
        y = base_y - int((base_y-top_y)*pct/80)
        d.line((130, y, 1550, y), fill=grid, width=2); d.text((55, y-13), str(pct), font=font(22), fill="#777777")
    groups = [("平均分", 64.53, 56.00), ("达标率", 63.33, 53.33)]
    for gi, (label, a, b) in enumerate(groups):
        cx = 520 + gi*650
        for off, val, color in [(-95, a, blue), (35, b, gold)]:
            h = int((base_y-top_y)*val/80); x1 = cx+off
            d.rounded_rectangle((x1, base_y-h, x1+90, base_y), radius=8, fill=color)
            d.text((x1+45, base_y-h-38), f"{val:.2f}", font=bold(22), fill=ink, anchor="mm")
        d.text((cx, 645), label, font=bold(26), fill=ink, anchor="mm")
    d.rectangle((1190, 30, 1220, 55), fill=blue); d.text((1235, 28), "三四年级", font=font(22), fill=ink)
    d.rectangle((1375, 30, 1405, 55), fill=gold); d.text((1420, 28), "五六年级", font=font(22), fill=ink)
    im.save(CHARTS/"overall_compare.png")

    # Distribution grouped bars.
    im, d = canvas(); base_y, top_y = 610, 100
    d.text((85, 30), "成绩分布对比", font=bold(34), fill=ink)
    for pct in range(0, 51, 10):
        y = base_y - int((base_y-top_y)*pct/50)
        d.line((130, y, 1570, y), fill=grid, width=2); d.text((45, y-13), f"{pct}%", font=font(20), fill="#777777")
    for i, ((label,_,a), (_,_,b)) in enumerate(zip(data_34["dist"], data_56["dist"])):
        cx = 285+i*285
        for off, val, color in [(-58,a,blue),(8,b,gold)]:
            h = int((base_y-top_y)*val/50); x1=cx+off
            d.rectangle((x1,base_y-h,x1+55,base_y),fill=color)
            d.text((x1+27,base_y-h-25),f"{val:.1f}%",font=font(17),fill=ink,anchor="mm")
        d.text((cx, 645), label, font=font(20), fill=ink, anchor="mm")
    d.rectangle((1190, 35, 1220, 60), fill=blue); d.text((1235, 31), "三四年级", font=font(22), fill=ink)
    d.rectangle((1375, 35, 1405, 60), fill=gold); d.text((1420, 31), "五六年级", font=font(22), fill=ink)
    im.save(CHARTS/"distribution.png")

    # Horizontal primary-dimension charts.
    for dataset, filename, color in [(data_34,"primary_34.png",blue),(data_56,"primary_56.png",gold)]:
        im = Image.new("RGB", (1640, 850), "white"); d = ImageDraw.Draw(im)
        d.text((70, 25), f"{dataset['name']}一级知识维度得分率", font=bold(34), fill=ink)
        left, right, top, row_h = 560, 1510, 115, 108
        ref_x = left + int((right-left)*.6)
        d.line((ref_x, top-10, ref_x, top+row_h*6-20), fill=red, width=3)
        d.text((ref_x+8, top-42), "60%参考线", font=font(19), fill=red)
        order = sorted(dataset["primary"], key=lambda x:x[3])
        for i,(label,_,_,val) in enumerate(order):
            y = top+i*row_h
            d.text((530,y+22),label,font=font(23),fill=ink,anchor="ra")
            d.rounded_rectangle((left,y+10,right,y+58),radius=8,fill="#EDF0F4")
            bar_right = left+int((right-left)*val/100)
            d.rounded_rectangle((left,y+10,bar_right,y+58),radius=8,fill=color)
            d.text((bar_right+14,y+34),f"{val:.2f}%",font=bold(21),fill=ink,anchor="lm")
        im.save(CHARTS/filename)


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Inches(.72); sec.bottom_margin = Inches(.68)
    sec.left_margin = Inches(1); sec.right_margin = Inches(1)
    sec.header_distance = Inches(.3); sec.footer_distance = Inches(.32)
    sec.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5); normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
    for i, size, color, before, after in [(1,16,NAVY,16,8),(2,13,BLUE,12,6),(3,11.5,TEAL,8,4)]:
        st = styles[f"Heading {i}"]
        st.font.name = "Arial"; st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for name in ["List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = "Arial"; st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(10.5); st.paragraph_format.left_indent = Inches(.5)
        st.paragraph_format.first_line_indent = Inches(-.25); st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.2

    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(hp, after=0, line=1.0)
    add_text(hp, "学生素养测评数据分析报告", size=8.5, color=GRAY)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(fp, after=0, line=1.0)
    add_text(fp, "学生素养测评数据分析报告  ·  ", size=8, color=GRAY)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)
    return doc


def add_cover(doc):
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_paragraph(p, after=8)
    add_text(p, "2026年", size=11, bold=True, color=GOLD)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_paragraph(p, after=10, line=1.1)
    add_text(p, "XXX学校", size=24, bold=True, color=NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_paragraph(p, after=8)
    add_text(p, "学生素养测评数据分析报告", size=29, bold=True, color=NAVY)
    doc.add_page_break()


def add_overview(doc):
    add_heading(doc, "一、学校整体情况", 1)
    add_callout(doc, "整体表现", "本次测评共纳入60份有效记录，学校整体平均分为60.27分，达标35人，达标率为58.33%。三四年级达标率为63.33%，五六年级达标率为53.33%。", fill=PALE_GOLD, accent=RED)

    add_heading(doc, "1. 整体成绩概况", 2)
    add_body(doc, "从整体成绩看，学校超过半数学生达到60分及以上。三四年级平均分为64.53分，五六年级平均分为56.00分；两个年级段在平均分、达标率和高分段人数方面呈现出不同特点。")
    add_table(doc, ["指标", "学校整体", "三四年级", "五六年级"], [
        ("有效样本", "60", "30", "30"),
        ("平均分", "60.27", "64.53", "56.00"),
        ("达标人数（≥60分）", "35", "19", "16"),
        ("达标率", "58.33%", "63.33%", "53.33%"),
        ("优秀段人数（≥90分）", "4", "4", "0"),
        ("优秀段占比", "6.67%", "13.33%", "0.00%"),
    ], [2600, 2260, 2250, 2250], font=9.5)
    add_picture(doc, CHARTS/"overall_compare.png", caption="图1  两个年龄段平均分与达标率对比")

    add_heading(doc, "2. 成绩分布与年级段特点", 2)
    combined = []
    for i, label in enumerate([x[0] for x in data_34["dist"]]):
        count = data_34["dist"][i][1] + data_56["dist"][i][1]
        combined.append((label, count, f"{count/60*100:.2f}%", data_34["dist"][i][1], data_56["dist"][i][1]))
    add_table(doc, ["分数段", "学校人数", "学校占比", "三四年级", "五六年级"], combined,
              [2100, 1600, 1800, 1930, 1930], font=9.2)
    add_picture(doc, CHARTS/"distribution.png", caption="图2  两个年龄段成绩分布对比")
    add_bullet(doc, "学校整体：60分以下25人，占41.67%；70分及以上22人，占36.67%。")
    add_bullet(doc, "三四年级：达标19人，80分及以上10人；高分学生占有一定比例，同时仍有11人低于60分。")
    add_bullet(doc, "五六年级：达标16人，60分以下14人；70分以下共22人，本次测评中没有学生进入90分及以上分数段。")

    add_heading(doc, "3. 各能力维度主要表现", 2)
    add_body(doc, "从一级知识维度看，两个年级段在数据、算法与模型方面表现相对较好；AI关键技术与应用形态以及AI伦理、安全与社会责任是后续需要重点关注的方向。")
    add_table(doc, ["观察方向", "三四年级表现", "五六年级表现", "学校层面判断"], [
        ("数据、算法与模型", "72.67%，表现较好", "62.78%，表现较好", "两个年级段的共同优势"),
        ("AI基础认知与智能社会", "58.33%", "64.67%，组内最高", "五六年级表现相对突出"),
        ("AI关键技术与应用形态", "55.83%", "41.11%，组内最低", "需要重点关注"),
        ("伦理、安全与社会责任", "51.90%，组内最低", "54.67%", "两个年级段均有提升空间"),
        ("系统设计与工程实践", "63.33%", "50.83%", "三四年级表现相对更好"),
    ], [2100, 1800, 1800, 3660], font=8.8, aligns=[WD_ALIGN_PARAGRAPH.LEFT]*4)


def age_section(doc, dataset, part_no, chart_file):
    name = dataset["name"]
    add_heading(doc, f"{part_no}、{name}测评情况", 1)
    pass_count = dataset["n"] - dataset["dist"][0][1]
    pass_rate = pass_count / dataset["n"] * 100
    primary_sorted = sorted(dataset["primary"], key=lambda x: x[3], reverse=True)
    secondary_sorted = sorted(dataset["secondary"], key=lambda x: x[1], reverse=True)

    add_heading(doc, "1. 关键结论摘要", 2)
    add_bullet(doc, f"本次纳入有效记录{dataset['n']}份，平均分为{dataset['mean']:.2f}分，达标率为{pass_rate:.2f}%。")
    add_bullet(doc, f"60分以下共{dataset['dist'][0][1]}人，占{dataset['dist'][0][2]:.2f}%，是人数最多的分数段。")
    add_bullet(doc, f"知识表现中，{primary_sorted[0][0]}相对较高，{primary_sorted[-1][0]}相对较低，二者得分率相差{primary_sorted[0][3]-primary_sorted[-1][3]:.2f}个百分点。")

    add_heading(doc, "2. 数据口径与样本概况", 2)
    add_body(doc, f"本次统计纳入{dataset['n']}份有效记录，测评成绩满分为100分。")

    add_heading(doc, "3. 整体表现分析", 2)
    add_body(doc, f"平均分为{dataset['mean']:.2f}分，中位数为{dataset['median']:.2f}分，标准差为{dataset['sd']:.2f}分。")
    add_table(doc, ["有效样本", "平均分", "中位数", "标准差", "最低分", "最高分"],
              [(dataset["n"], f"{dataset['mean']:.2f}", f"{dataset['median']:.2f}", f"{dataset['sd']:.2f}", f"{dataset['min']:.2f}", f"{dataset['max']:.2f}")],
              [1560]*6, font=9.2)
    add_table(doc, ["分数段", "人数", "占比"], [(a,b,f"{c:.2f}%") for a,b,c in dataset["dist"]], [4200, 2400, 2760], font=9.5)

    add_heading(doc, "4. 一级知识维度分析", 2)
    primary_avg = 62.23 if name == "三四年级" else 54.01
    add_body(doc, f"本板块共覆盖6个维度，平均得分率为{primary_avg:.2f}%。{primary_sorted[0][0]}相对较高（{primary_sorted[0][3]:.2f}%），{primary_sorted[-1][0]}相对较低（{primary_sorted[-1][3]:.2f}%）。")
    add_picture(doc, chart_file, caption=f"图  {name}一级知识维度得分率")
    rows = [(n, q, f"{m:.0f}", f"{r:.2f}%") for n,q,m,r in dataset["primary"]]
    add_table(doc, ["一级知识维度", "题目数", "满分", "得分率"], rows,
              [4800, 1320, 1320, 1920], font=8.9)

    add_heading(doc, "5. 二级知识维度分析", 2)
    secondary_avg = 62.05 if name == "三四年级" else 56.00
    add_body(doc, f"本板块共覆盖{len(dataset['secondary'])}个维度，平均得分率为{secondary_avg:.2f}%。{secondary_sorted[0][0]}相对较高（{secondary_sorted[0][1]:.2f}%），{secondary_sorted[-1][0]}相对较低（{secondary_sorted[-1][1]:.2f}%）。")
    half = (len(dataset["secondary"]) + 1) // 2
    secondary_rows = []
    for i in range(half):
        left_name, left_rate = dataset["secondary"][i]
        if i + half < len(dataset["secondary"]):
            right_name, right_rate = dataset["secondary"][i + half]
            secondary_rows.append((left_name, f"{left_rate:.2f}%", right_name, f"{right_rate:.2f}%"))
        else:
            secondary_rows.append((left_name, f"{left_rate:.2f}%", "", ""))
    add_table(doc, ["二级知识维度", "得分率", "二级知识维度", "得分率"],
              secondary_rows, [3300, 1380, 3300, 1380], font=8.3)

    add_heading(doc, "6. 重点题目诊断", 2)
    avg_low = sum(x[3] for x in dataset["items"]) / len(dataset["items"])
    add_body(doc, f"按相对得分率选取5道重点复核题，平均得分率为{avg_low:.2f}%。建议结合题干语言、选项干扰、作答方式和学生认知负荷逐题复核。")
    add_table(doc, ["题号", "一级知识维度", "二级知识维度", "得分率"],
              [(x[0], x[1], x[2], f"{x[3]:.2f}%") for x in dataset["items"]], [900, 3300, 3300, 1860], font=8.8)


def main():
    make_charts()
    doc = setup_document()
    add_cover(doc)
    add_overview(doc)
    age_section(doc, data_34, "二", CHARTS/"primary_34.png")
    age_section(doc, data_56, "三", CHARTS/"primary_56.png")
    props = doc.core_properties
    props.title = "XXX学校学生素养测评数据分析报告"
    props.subject = "学生素养测评数据分析"
    props.author = ""
    props.keywords = "素养测评, 数据分析, 学校汇总"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
